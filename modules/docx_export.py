"""逐語録の Word (.docx) エクスポートモジュール

sessions.transcript_json（AssemblyAIのutterancesそのまま）から、
ユーザーが自分で編集できる素材としての逐語録docxを生成する。
話者識別などの精度は保証しない位置づけ（引き継ぎ書§2-D）。
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt

# TODO: フォーマットは暫定（引き継ぎ書§4）: 話者ラベルあり・タイムスタンプなし。
# transcript_json には start/end (ms) が保存されているので、タイムスタンプ付きが
# 必要になったらここで出し分けられる。

_NOTICE = (
    "本逐語録はAIによる自動文字起こしです。話者の識別や語句に誤りを含む可能性があります。"
    "クライアントの同意の範囲内でご利用いただき、第三者への共有・二次配布にはご注意ください。"
)


def generate_transcript_docx(transcript_json: list[dict], created_at: datetime) -> bytes:
    """話者ラベル付き逐語録のdocxバイト列を生成する"""
    doc = Document()

    doc.add_heading("コーチングセッション逐語録", level=1)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"分析日: {created_at.strftime('%Y年%m月%d日')}")
    meta_run.font.size = Pt(10)

    notice = doc.add_paragraph()
    notice_run = notice.add_run(_NOTICE)
    notice_run.font.size = Pt(9)
    notice_run.italic = True

    doc.add_paragraph()  # 区切りの空行

    for utt in transcript_json or []:
        speaker = utt.get("speaker") or "不明"
        text = utt.get("text") or ""
        p = doc.add_paragraph()
        speaker_run = p.add_run(f"{speaker}：")
        speaker_run.bold = True
        p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
